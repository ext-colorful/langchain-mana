"""File upload and parsing service."""

from __future__ import annotations

import os
from pathlib import Path
from typing import Tuple

import aiofiles
from fastapi import UploadFile
from pypdf import PdfReader
from docx import Document as DocxDocument

from app.core.config import settings
from app.core.exceptions import FileProcessingError
from app.core.logging import logger

ALLOWED_EXTENSIONS = set(settings.UPLOAD_ALLOWED_EXTENSIONS)


class FileService:
    """Handles file uploads and parsing for RAG."""

    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    def _validate_file(self, filename: str) -> None:
        ext = os.path.splitext(filename)[1].lower()
        if ext not in ALLOWED_EXTENSIONS:
            raise FileProcessingError(f"Unsupported file type: {ext}")

    async def save_file(self, upload_file: UploadFile) -> Path:
        self._validate_file(upload_file.filename)
        file_path = self.upload_dir / upload_file.filename

        async with aiofiles.open(file_path, "wb") as out_file:
            while content := await upload_file.read(1024 * 1024):
                await out_file.write(content)

        logger.info(f"File saved: {file_path}")
        return file_path

    def parse_file(self, file_path: Path) -> Tuple[str, dict]:
        ext = file_path.suffix.lower()
        if ext == ".pdf":
            return self._parse_pdf(file_path)
        if ext == ".docx":
            return self._parse_docx(file_path)
        if ext in {".txt", ".md"}:
            text = file_path.read_text(encoding="utf-8")
            return text, {"type": "text"}
        raise FileProcessingError(f"Unsupported file type: {ext}")

    def _parse_pdf(self, file_path: Path) -> Tuple[str, dict]:
        reader = PdfReader(str(file_path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text, {"type": "pdf", "pages": len(reader.pages)}

    def _parse_docx(self, file_path: Path) -> Tuple[str, dict]:
        doc = DocxDocument(str(file_path))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        return text, {"type": "docx", "paragraphs": len(doc.paragraphs)}
